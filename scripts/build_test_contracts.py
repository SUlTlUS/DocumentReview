from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "backend" / "tests" / "fixtures"

TITLE = "采购服务合同"
PARTIES = [
    "甲方：示例科技有限公司",
    "乙方：示例供应商有限公司",
]
SECTIONS = [
    ("一、合同标的", "乙方向甲方提供办公设备及配套安装服务，合同总价为人民币壹拾万元整。"),
    ("二、交付与验收", "乙方应在合同生效后十个工作日内交付。甲方应在合理期限内完成验收，如未提出异议视为验收合格。"),
    ("三、付款", "甲方在验收合格并收到合法发票后十五个工作日内付款。"),
    ("四、违约责任", "乙方逾期交付的，每日按合同总价的千分之一支付违约金；违约责任仅约束乙方。"),
    ("五、争议解决", "双方发生争议时应先协商，协商不成提交上海仲裁委员会仲裁。"),
]


def set_run_font(run, name: str, size: float, color: str = "000000", bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def build_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run(TITLE)
    set_run_font(title_run, "Microsoft YaHei", 22, "0B2545", True)

    for party in PARTIES:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(party)
        set_run_font(run, "Microsoft YaHei", 11, "333333")

    document.add_paragraph()
    for heading, body in SECTIONS:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.paragraph_format.space_before = Pt(11)
        heading_paragraph.paragraph_format.space_after = Pt(6)
        heading_run = heading_paragraph.add_run(heading)
        set_run_font(heading_run, "Microsoft YaHei", 13, "2E74B5", True)
        body_paragraph = document.add_paragraph()
        body_paragraph.paragraph_format.line_spacing = 1.25
        body_run = body_paragraph.add_run(body)
        set_run_font(body_run, "Microsoft YaHei", 11, "222222")

    document.core_properties.title = TITLE
    document.core_properties.subject = "RAG 文档审核测试样例"
    document.core_properties.author = "RAG Document Review"
    document.save(path)


def build_pdf(path: Path) -> None:
    font_path = Path("C:/Windows/Fonts/Deng.ttf")
    if not font_path.is_file():
        raise FileNotFoundError(f"Required CJK test font not found: {font_path}")
    pdfmetrics.registerFont(TTFont("DengXian", str(font_path)))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ContractTitle",
        parent=styles["Title"],
        fontName="DengXian",
        fontSize=22,
        leading=28,
        alignment=TA_CENTER,
        textColor=HexColor("#0B2545"),
        spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "ContractBody",
        parent=styles["BodyText"],
        fontName="DengXian",
        fontSize=11,
        leading=17,
        alignment=TA_LEFT,
        textColor=HexColor("#222222"),
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "ContractHeading",
        parent=body_style,
        fontSize=13,
        leading=19,
        textColor=HexColor("#2E74B5"),
        spaceBefore=11,
        spaceAfter=6,
    )
    story = [Paragraph(TITLE, title_style)]
    story.extend(Paragraph(party, body_style) for party in PARTIES)
    story.append(Spacer(1, 0.12 * inch))
    for heading, body in SECTIONS:
        story.append(Paragraph(heading, heading_style))
        story.append(Paragraph(body, body_style))
    document = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=TITLE,
        author="RAG Document Review",
    )
    document.build(story)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(OUTPUT_DIR / "test_contract.docx")
    build_pdf(OUTPUT_DIR / "test_contract.pdf")
    print("Generated test_contract.docx and test_contract.pdf")


if __name__ == "__main__":
    main()
