from dataclasses import dataclass
import logging
from pathlib import Path

from docx import Document as DocxDocument
from PyPDF2 import PdfReader

from app.errors import ParseError
from app.engines.ocr import OCRProvider


SUPPORTED_TYPES = {"pdf", "docx", "txt"}
logger = logging.getLogger("uvicorn.error")


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    method: str


def _read_txt(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ParseError("TXT 文件编码无法识别", "请将文件另存为 UTF-8 后重试")


def _read_docx(path: Path) -> str:
    document = DocxDocument(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)


def parse_document(
    file_path: str | Path,
    file_type: str,
    ocr_provider: OCRProvider | None = None,
) -> ParsedDocument:
    path = Path(file_path)
    normalized_type = file_type.lower().lstrip(".")
    if normalized_type not in SUPPORTED_TYPES:
        raise ParseError(f"不支持的文件类型：{file_type}")
    if not path.is_file():
        raise ParseError("上传文件不存在")

    try:
        if normalized_type == "txt":
            text = _read_txt(path)
        elif normalized_type == "docx":
            text = _read_docx(path)
        elif normalized_type == "pdf":
            text = _read_pdf(path)
            if not text.strip() and ocr_provider is not None:
                logger.info("scan_pdf_detected ocr_provider=%s", ocr_provider.name)
                text = ocr_provider.process_pdf(path)
                method = ocr_provider.name
            else:
                method = "digital"
        else:
            raise ParseError(f"不支持的文件类型：{file_type}")
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError("文档解析失败", str(exc)) from exc

    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        raise ParseError("文档中未提取到可审核文字")
    return ParsedDocument(text=normalized, method=locals().get("method", "digital"))


def parse_file(
    file_path: str | Path,
    file_type: str,
    ocr_provider: OCRProvider | None = None,
) -> str:
    return parse_document(file_path, file_type, ocr_provider).text
