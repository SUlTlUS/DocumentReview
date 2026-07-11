from pathlib import Path

from docx import Document as DocxDocument

from app.services.parser import parse_file
from app.services.rag import RAGRegistry, RAGService, split_text


FIXTURES = Path(__file__).parent / "fixtures"


def test_parse_txt_supports_utf8_and_gb18030(tmp_path):
    utf8_path = tmp_path / "utf8.txt"
    utf8_path.write_text("违约责任：逾期付款应承担违约金。", encoding="utf-8")
    gb_path = tmp_path / "gb.txt"
    gb_path.write_bytes("争议解决：提交仲裁。".encode("gb18030"))

    assert "违约责任" in parse_file(utf8_path, "txt")
    assert "争议解决" in parse_file(gb_path, ".txt")


def test_parse_docx_extracts_paragraphs(tmp_path):
    path = tmp_path / "contract.docx"
    document = DocxDocument()
    document.add_paragraph("采购合同")
    document.add_paragraph("甲方应在三十日内验收。")
    document.save(path)

    assert parse_file(path, "docx") == "采购合同\n甲方应在三十日内验收。"


def test_real_contract_fixtures_extract_pdf_docx_and_txt():
    for extension in ("txt", "docx", "pdf"):
        extracted = parse_file(FIXTURES / f"test_contract.{extension}", extension)
        assert "采购服务合同" in extracted
        assert "合理期限" in extracted
        assert "违约责任仅约束乙方" in extracted


def test_split_text_has_overlap_and_respects_size():
    text = "甲" * 450 + "。" + "违约责任" * 80
    chunks = split_text(text, chunk_size=500, overlap=100)

    assert len(chunks) >= 2
    assert all(len(chunk) <= 500 for chunk in chunks)
    assert set(chunks[0][-80:]) & set(chunks[1][:120])


def test_rag_returns_relevant_chunk_first():
    rag = RAGService(chunk_size=35, overlap=5)
    rag.build_index(
        "交付期限为十个工作日。\n"
        "违约责任：逾期付款按日万分之五支付违约金。\n"
        "争议由上海仲裁委员会处理。"
    )

    results = rag.retrieve_context("逾期违约金", top_k=2)
    assert "违约" in results[0].content
    assert results[0].score >= results[1].score


def test_registry_isolates_documents_and_rebuilds_after_restart():
    registry = RAGRegistry()
    registry.build(1, "保密义务持续三年。")
    registry.build(2, "付款期限为收到发票后十五日。")

    assert "保密" in registry.get_or_build(1, "unused").get_chunks()[0]
    registry.clear()
    rebuilt = registry.get_or_build(1, "保密义务持续三年。")
    assert "保密" in rebuilt.retrieve_context("保密")[0].content
